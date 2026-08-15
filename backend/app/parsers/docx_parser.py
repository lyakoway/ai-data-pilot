"""Word (.docx) parsing — pseudo-pages by char budget."""
from __future__ import annotations

from pathlib import Path

from app.parsers.base import PageSegment, ParseError

_PAGE_CHAR_BUDGET = 1800


def parse(path: Path) -> list[PageSegment]:
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:  # noqa: BLE001
        raise ParseError(f"Не удалось прочитать Word-документ: {exc}") from exc

    blocks: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    if not blocks:
        raise ParseError("Word-документ пуст или не содержит текста.")

    segments: list[PageSegment] = []
    buf: list[str] = []
    size = 0
    page = 1
    for block in blocks:
        buf.append(block)
        size += len(block)
        if size >= _PAGE_CHAR_BUDGET:
            segments.append(PageSegment(page=page, text="\n".join(buf), label=f"раздел {page}"))
            buf, size, page = [], 0, page + 1
    if buf:
        segments.append(PageSegment(page=page, text="\n".join(buf), label=f"раздел {page}"))
    return segments
